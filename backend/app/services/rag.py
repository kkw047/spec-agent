from pathlib import Path
import re
from typing import Iterable

import fitz
import psycopg
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_postgres import PGVector
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.core.config import Settings
from app.models.schemas import ReferenceItem


NOISY_PATENT_PAGE_PHRASES = [
    "JavaScript 를 사용할 수 없습니다",
    "본문내용 바로가기",
    "사이트맵",
    "화면크기",
    "특허고객등록",
    "전자출원 상담지원",
]

USEFUL_GUIDE_MARKERS = [
    "【발명의 명칭】",
    "【기술분야】",
    "【발명의 배경이 되는 기술】",
    "【해결하려는 과제】",
    "【과제의 해결 수단】",
    "【발명의 효과】",
    "【발명을 실시하기 위한 구체적인 내용】",
    "【산업상 이용가능성】",
    "【청구범위】",
]

REFERENCE_TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".log"}


#  pgVector 준비
def ensure_vector_extension(settings: Settings) -> None:
    if not settings.postgres_user:
        return
    with psycopg.connect(**settings.psycopg_params) as conn:
        with conn.cursor() as cursor:
            cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
        conn.commit()


# 세션별 컬렉션 이름
# 사용자 사건 자료는 공용 참고자료와 섞이지 않도록 session_id별 컬렉션에 저장
def case_collection_name(settings: Settings, session_id: str) -> str:
    safe_id = "".join(ch for ch in session_id if ch.isalnum() or ch in ("_", "-"))[:80]
    return f"{settings.pgvector_collection}_case_{safe_id}"


# 참고자료 PDF 읽기
# PDF 텍스트를 읽고, 글자가 깨지면 OCR로 다시 시도
def load_pdf_documents(pdf_path: Path) -> list[Document]:
    reader = PdfReader(str(pdf_path))
    documents = []
    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": str(pdf_path), "page": index, "title": pdf_path.name},
                )
            )

    if documents and _garbled_ratio(documents) > 0.35:
        return _ocr_pdf(pdf_path)
    return documents


# 참고자료 텍스트 읽기
# local_data/references의 TXT/MD/CSV/JSON/LOG 파일을 Document로 제작
def load_text_documents(path: Path) -> list[Document]:
    text = _read_text_file(path)
    if not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": str(path), "title": path.name, "kind": "text"},
        )
    ]


#  참고자료 DOCX 읽기
# 문단 텍스트를 공용 RAG 자료로 사용
def load_docx_documents(path: Path) -> list[Document]:
    document = DocxDocument(str(path))
    text = "\n".join(
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    )
    if not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": str(path), "title": path.name, "kind": "docx"},
        )
    ]


# 참고자료 파일 라우팅
# 파일 확장자에 따라 PDF/TXT/DOCX 로더를 선택
def load_reference_file_documents(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf_documents(path)
    if suffix == ".docx":
        return load_docx_documents(path)
    if suffix in REFERENCE_TEXT_SUFFIXES and not path.name.endswith(".ocr.txt"):
        return load_text_documents(path)
    return []


#  인코딩 처리
# 한국어 문서가 UTF-8이 아닐 수 있어 cp949/euc-kr까지 시도합니다.
def _read_text_file(path: Path) -> str:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


# PDF 깨짐 판단
# 깨진 텍스트 비율이 높으면 OCR 경로로 전환합니다.
def _garbled_ratio(documents: list[Document]) -> float:
    if not documents:
        return 0
    return sum(1 for doc in documents if _looks_garbled(doc.page_content)) / len(documents)


def _looks_garbled(text: str) -> bool:
    stripped = text.strip()
    if len(stripped) < 80:
        return False
    readable_count = sum(
        1
        for ch in stripped
        if ch.isspace()
        or ch.isascii()
        or "\uac00" <= ch <= "\ud7a3"
        or ch in "ㆍ·,.;:()[]{}+-/%"
    )
    return readable_count / len(stripped) < 0.72


#  PDF OCR
# OCR 결과는 같은 위치의 *.ocr.txt에 캐시
def _ocr_pdf(pdf_path: Path) -> list[Document]:
    cache_path = pdf_path.with_suffix(".ocr.txt")
    if cache_path.exists():
        return _read_ocr_cache(cache_path, pdf_path)

    pages = []
    with fitz.open(pdf_path) as doc:
        for index, page in enumerate(doc, 1):
            textpage = page.get_textpage_ocr(language="kor+eng", dpi=200, full=True)
            text = page.get_text(textpage=textpage).strip()
            if text:
                pages.append((index, text))
    _write_ocr_cache(cache_path, pages)
    return [
        Document(
            page_content=text,
            metadata={"source": str(pdf_path), "page": index, "title": pdf_path.name},
        )
        for index, text in pages
    ]


def _write_ocr_cache(cache_path: Path, pages: list[tuple[int, str]]) -> None:
    content = "\n\n".join(f"===== PAGE {index} =====\n{text}" for index, text in pages)
    cache_path.write_text(content, encoding="utf-8")


def _read_ocr_cache(cache_path: Path, pdf_path: Path) -> list[Document]:
    raw = cache_path.read_text(encoding="utf-8")
    documents = []
    for chunk in raw.split("===== PAGE "):
        chunk = chunk.strip()
        if not chunk:
            continue
        header, _, text = chunk.partition("=====")
        try:
            index = int(header.strip())
        except ValueError:
            continue
        documents.append(
            Document(
                page_content=text.strip(),
                metadata={"source": str(pdf_path), "page": index, "title": pdf_path.name},
            )
        )
    return documents


# 특허로 안내 페이지 수집
# requests.get(url)로 특허로 명세서 작성 안내 HTML을 가져와 공용 참고자료로 제작
def load_patent_guide(url: str) -> list[Document]:
    response = requests.get(url, timeout=20)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
    text = _clean_patent_guide_text(text)
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": url, "title": "특허로 출원명세서 안내"})]


# 특허로 페이지 정리
# 메뉴, 다운로드 목록, JavaScript 안내처럼 RAG에 방해되는 문장을 제거합니다.
def _clean_patent_guide_text(text: str) -> str:
    lines = []
    for line in text.splitlines():
        compact = line.strip()
        if not compact:
            continue
        if _looks_like_download_catalog(compact):
            continue
        if any(phrase in compact for phrase in NOISY_PATENT_PAGE_PHRASES):
            continue
        lines.append(compact)
    cleaned = "\n".join(lines)
    marker_positions = [cleaned.find(marker) for marker in USEFUL_GUIDE_MARKERS if marker in cleaned]
    if marker_positions:
        cleaned = cleaned[min(marker_positions) :]
    return cleaned


def _looks_like_download_catalog(text: str) -> bool:
    if text.count("다운로드") >= 2:
        return True
    if "다운로드" in text and len(text) > 90:
        return True
    return False


#  토큰화
# 긴 문서를 900자 단위, 120자 겹침 청크로 나눕니다.
def split_documents(documents: Iterable[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=120)
    return splitter.split_documents(list(documents))


# 공용 참고자료 DB 저장
# local_data/references + 특허로 안내 자료를 공용 pgVector 컬렉션에 저장합니다.
def build_vectorstore(settings: Settings, documents: list[Document], reset: bool = False) -> int:
    return add_documents_to_collection(
        settings=settings,
        documents=documents,
        collection_name=settings.pgvector_collection,
        reset=reset,
    )


# 임베딩 -> pgVector 저장
# OpenAIEmbeddings로 청크 벡터를 만들고 PGVector.from_documents()로 DB에 저장
def add_documents_to_collection(
    settings: Settings,
    documents: list[Document],
    collection_name: str,
    reset: bool = False,
) -> int:
    if not documents or not settings.sqlalchemy_database_url or not settings.openai_api_key:
        return 0
    ensure_vector_extension(settings)
    embeddings = OpenAIEmbeddings(
        model=settings.openai_embedding_model,
        api_key=settings.openai_api_key,
    )
    PGVector.from_documents(
        documents=documents,
        embedding=embeddings,
        collection_name=collection_name,
        connection=settings.sqlalchemy_database_url,
        pre_delete_collection=reset,
    )
    return len(documents)


#  공용 참고자료 검색
# 사전에 인덱싱한 특허로/수업자료/참고 PDF와 query를 비교
def get_references(settings: Settings, query: str, k: int = 4) -> list[ReferenceItem]:
    return get_references_from_collection(settings, settings.pgvector_collection, query, k=k)


# 현재 사건 자료 검색
# 사용자가 이번 세션에 올린 파일/메시지 청크와 query를 비교
def get_case_references(settings: Settings, session_id: str, query: str, k: int = 6) -> list[ReferenceItem]:
    return get_references_from_collection(settings, case_collection_name(settings, session_id), query, k=k)


# similarity_search 실행
# 검색 대상은 인터넷이 아니라 pgVector 컬렉션에 이미 저장된 청크
def get_references_from_collection(
    settings: Settings,
    collection_name: str,
    query: str,
    k: int = 4,
) -> list[ReferenceItem]:
    if not settings.sqlalchemy_database_url or not settings.openai_api_key or not query.strip():
        return []
    try:
        # DB 검색 호출
        # query -> OpenAI embedding -> pgVector similarity_search 순서로 비교
        embeddings = OpenAIEmbeddings(
            model=settings.openai_embedding_model,
            api_key=settings.openai_api_key,
        )
        vectorstore = PGVector(
            embeddings=embeddings,
            collection_name=collection_name,
            connection=settings.sqlalchemy_database_url,
        )
        docs = vectorstore.similarity_search(query, k=k)
    except Exception:
        return []

    references = []
    for doc in docs:
        title = str(doc.metadata.get("title") or "참고자료")
        if title == "사용자 메시지":
            continue
        excerpt = clean_reference_excerpt(doc.page_content)
        if _is_noisy_reference(excerpt, doc.metadata):
            continue
        references.append(
            ReferenceItem(
                title=title,
                source=str(doc.metadata.get("source") or collection_name),
                excerpt=excerpt[:700],
            )
        )
    return references


def clean_reference_excerpt(text: str) -> str:
    cleaned = " ".join((text or "").split())
    marker_positions = [cleaned.find(marker) for marker in USEFUL_GUIDE_MARKERS if marker in cleaned]
    if marker_positions:
        cleaned = cleaned[min(marker_positions) :]
    cleaned = re.sub(r"특허로\s+JavaScript\s+를 사용할 수 없습니다.*?본문내용 바로가기", "", cleaned)
    return cleaned.strip()


def _is_noisy_reference(excerpt: str, metadata: dict) -> bool:
    if not excerpt:
        return True
    title = str(metadata.get("title") or "")
    source = str(metadata.get("source") or "")
    is_patent_guide = "특허로 출원명세서 안내" in title or "patent.go.kr" in source
    if not is_patent_guide:
        return False
    if any(marker in excerpt for marker in USEFUL_GUIDE_MARKERS):
        return False
    if _looks_like_download_catalog(excerpt):
        return True
    return any(phrase in excerpt for phrase in NOISY_PATENT_PAGE_PHRASES)
