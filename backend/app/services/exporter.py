from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt

from app.models.schemas import DraftResponse
from app.services.markdown import build_markdown


TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "templates" / "명세서_양식.docx"
XML_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
KOREAN_FONT = "맑은 고딕"
REVIEW_RED = RGBColor(0xC0, 0x00, 0x00)
HUMAN_REVIEW_MARKERS = [
    "작성 필요",
    "확인 필요",
    "검토 필요",
    "사람 검토",
    "변리사 검토",
    "최종 특허성 판단",
    "청구범위 확정",
    "출원 여부",
    "도면 파일",
    "도면 설명",
    "부호",
    "부족",
    "차단",
    "주의",
]


# Word XML에 들어갈 수 없는 제어문자 제거
def _xml_safe(text: str) -> str:
    return XML_CONTROL_CHARS.sub("", str(text or ""))


# 사람 검수 표시 판단
# Word에서 빨간색으로 보여줄 문구인지 판단함.
def _needs_human_review(text: str) -> bool:
    normalized = _xml_safe(text)
    return any(marker in normalized for marker in HUMAN_REVIEW_MARKERS)


# run 색상 지정
# 사람이 봐야 하는 문구는 빨간색으로 표시함.
def _style_run(run, red: bool = False, size=None, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = KOREAN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        run.font.size = size
    if red:
        run.font.color.rgb = REVIEW_RED


# 문서 기본 한글 폰트
# Word에서 한글이 기본 서체 문제로 깨져 보이는 일을 줄임.
def _set_document_font(document: Document) -> None:
    for style_name in ["Normal", "Body Text"]:
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = KOREAN_FONT
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), KOREAN_FONT)


#  파일명 생성
# 산출물 이름은 명세서_양식_{사건명}.md/docx 형식으로 저장
def _safe_name(name: str) -> str:
    allowed = []
    for ch in name.strip():
        if ch.isalnum() or ch in ("-", "_"):
            allowed.append(ch)
        elif ch.isspace():
            allowed.append("_")
    return "".join(allowed)[:80] or "spec_agent_draft"


# 템플릿 채우기
# 템플릿의 【발명의 명칭】 같은 heading 바로 아래 빈 문단에 내용 삽입
def _set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    lines = (_xml_safe(text) or "작성 필요").splitlines() or ["작성 필요"]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _style_run(run, red=_needs_human_review(line))


def _fill_after_heading(document: Document, heading: str, text: str) -> None:
    paragraphs = list(document.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != heading:
            continue
        for candidate in paragraphs[index + 1 :]:
            if not candidate.text.strip():
                _set_paragraph_text(candidate, text)
                return
        paragraph = document.add_paragraph()
        _set_paragraph_text(paragraph, text)
        return


#  업로드 이미지 선택
# Word 대표도/도 1에 넣을 이미지 자료 선택
def _image_materials(response: DraftResponse) -> list:
    return [
        material
        for material in getattr(response, "materials", [])
        if material.kind == "image" and material.stored_path
    ]


# 대표도 이미지 삽입
# 이미지는 도면 후보로만 삽입하고, 구성요소/부호 해석은 자동 확정 X
def _fill_image_after_heading(document: Document, heading: str, material, fallback_text: str) -> bool:
    paragraphs = list(document.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != heading:
            continue
        target = None
        for candidate in paragraphs[index + 1 :]:
            if not candidate.text.strip():
                target = candidate
                break
        if target is None:
            target = document.add_paragraph()
        target.clear()
        image_run = target.add_run(_xml_safe(f"{material.name} (업로드 이미지, 도면 후보, 사람 검토 필요)"))
        _style_run(image_run, red=True)
        target.add_run().add_break()
        try:
            target.add_run().add_picture(material.stored_path, width=Inches(4.8))
        except Exception:
            target.add_run().add_break()
            fallback_run = target.add_run(_xml_safe(fallback_text))
            _style_run(fallback_run, red=True)
        return True
    return False


def _summary(response: DraftResponse) -> str:
    parts = [
        response.sections.invention_title,
        response.sections.problem_to_solve,
        response.sections.solution,
        response.sections.advantageous_effects,
    ]
    body = " ".join(part.strip() for part in parts if part and part.strip())
    return body[:700] or "요약 작성 필요"


def _claim_draft(response: DraftResponse) -> str:
    title = response.sections.invention_title or "발명"
    solution = response.sections.solution or response.sections.embodiment
    if not solution:
        return "청구항 초안 작성에는 구성요소와 결합관계 보완이 필요합니다."
    return (
        "검토용 청구항 초안입니다. 최종 청구범위 확정은 변리사 검토가 필요합니다.\n"
        f"청구항 1. {title}에 있어서, {solution}"
    )


#  Markdown / Word 산출물 생성
# 다운로드 파일은 명세서 초안 본문만 저장하고, 검토표는 웹 화면에서만 보여줌.
def export_draft(response: DraftResponse, output_dir: Path) -> tuple[Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = _safe_name(response.case_name)
    markdown_path = output_dir / f"명세서_양식_{stem}.md"
    docx_path = output_dir / f"명세서_양식_{stem}.docx"

    markdown = build_markdown(response)
    markdown_path.write_text(markdown, encoding="utf-8")

    document = Document(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else Document()
    _set_document_font(document)
    if not TEMPLATE_PATH.exists():
        document.add_heading("【발명의 설명】", level=1)

    fill_map = {
        "【발명의 명칭】": response.sections.invention_title,
        "【기술분야】": response.sections.technical_field,
        "【발명의 배경이 되는 기술】": response.sections.background_art,
        "【해결하고자 하는 과제】": response.sections.problem_to_solve,
        "【과제의 해결 수단】": response.sections.solution,
        "【발명의 효과】": response.sections.advantageous_effects,
        "【도면의 간단한 설명】": response.sections.drawing_description,
        "【발명을 실시하기 위한 구체적인 내용】": response.sections.embodiment,
        "【부호의 설명】": response.sections.reference_signs,
        "【청구항 1】": _claim_draft(response),
        "【요약】": _summary(response),
    }
    for heading, body in fill_map.items():
        _fill_after_heading(document, heading, body)

    image_materials = _image_materials(response)
    if image_materials:
        first_image = image_materials[0]
        image_note = "이미지를 Word에 삽입하지 못했습니다. 업로드 원본 파일을 확인해 주세요."
        _fill_image_after_heading(document, "【대표도】", first_image, image_note)
        _fill_image_after_heading(document, "【도 1】", first_image, image_note)
    else:
        _fill_after_heading(
            document,
            "【대표도】",
            "도 1" if response.sections.drawing_description else "대표도 확인 필요",
        )
        _fill_after_heading(
            document,
            "【도 1】",
            response.sections.drawing_description or "도면 파일 또는 도면 설명 확인 필요",
        )

    document.save(docx_path)
    return markdown_path, docx_path
