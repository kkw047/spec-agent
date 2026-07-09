from io import BytesIO
from pathlib import Path
import subprocess
import tempfile

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from app.models.schemas import MaterialSource
from app.services.rag import split_documents

#파일 처리 관련파일
#TXT, MD, CSV, JSON, LOG, PDF, DOCX, PNG, JPG, JPEG, WEBP, BMP, TIF, TIFF 지원함

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".log"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


#  업로드 파일명 정리
# 저장 경로에 위험한 문자가 들어가지 않도록 안전한 이름으로 변경
def safe_filename(name: str) -> str:
    stem = Path(name or "material").stem
    suffix = Path(name or "").suffix.lower()
    allowed = []
    for ch in stem:
        if ch.isalnum() or ch in ("-", "_"):
            allowed.append(ch)
        elif ch.isspace():
            allowed.append("_")
    safe_stem = "".join(allowed)[:80] or "material"
    return f"{safe_stem}{suffix}"


# 파일 입력 처리
# filename/content -> local_data/sessions/{session}/uploads 저장 -> 텍스트 추출 -> Document 청크 생성.
def extract_upload(filename: str, content: bytes, upload_dir: Path) -> tuple[MaterialSource, list[Document], str]:
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_name = safe_filename(filename)
    saved_path = upload_dir / safe_name
    saved_path.write_bytes(content)

    suffix = saved_path.suffix.lower()
    note = ""
    try:
        # 파일 종류별 추출
        # PDF/DOCX/TXT는 텍스트를 읽고, 이미지는 도면 후보로 저장하며 OCR만 시도합니다.
        if suffix == ".pdf":
            text = _extract_pdf(content)
            kind = "pdf"
        elif suffix == ".docx":
            text = _extract_docx(content)
            kind = "docx"
        elif suffix in TEXT_SUFFIXES:
            text = _decode_text(content)
            kind = "text"
        elif suffix in IMAGE_SUFFIXES:
            kind = "image"
            ocr_text = _extract_image(saved_path)
            text = _image_material_text(safe_name, ocr_text)
            if ocr_text.strip():
                note = "이미지 OCR 텍스트를 추출했습니다. 도면 구성과 부호는 사람이 확인해야 합니다."
            else:
                note = "이미지 파일을 도면 후보로 저장했습니다. 자동 시각 해석은 아직 지원하지 않으므로 도면 설명과 부호를 보완해 주세요."
        else:
            text = ""
            kind = "file"
            note = "지원하지 않는 파일 형식입니다. 핵심 내용을 메시지로 붙여 주세요."
    except Exception as exc:
        text = ""
        kind = suffix.lstrip(".") or "file"
        note = f"텍스트 추출 실패: {exc}"

    documents = []
    if text.strip():
        # Document 변환
        # split_documents()로 pgVector 저장용 청크 변환
        documents = [
            Document(
                page_content=text,
                metadata={
                    "source": str(saved_path),
                    "title": safe_name,
                    "kind": kind,
                },
            )
        ]
    chunks = split_documents(documents) if documents else []
    material = MaterialSource(
        name=safe_name,
        kind=kind,
        status="processed" if text.strip() else "needs_review",
        char_count=len(text),
        chunk_count=len(chunks),
        note=note,
        stored_path=str(saved_path),
    )
    return material, chunks, text


# 사용자 메시지 처리
# 채팅 입력도 파일과 동일하게 RAG/DB 저장 대상 Document 변환
def message_to_documents(message: str, session_id: str) -> list[Document]:
    if not message.strip():
        return []
    return split_documents(
        [
            Document(
                page_content=message,
                metadata={
                    "source": f"session:{session_id}:message",
                    "title": "사용자 메시지",
                    "kind": "message",
                },
            )
        ]
    )


# 텍스트 파일 디코딩
# 한국어 자료가 cp949/euc-kr일 수 있어 여러 인코딩을 순서대로 시도합니다.
def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


# PDF 텍스트 추출
# 업로드 PDF의 페이지별 텍스트를 읽어 하나의 문자열로 합침
def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[PDF page {index}]\n{text.strip()}")
    return "\n\n".join(pages)


# DOCX 텍스트 추출
# 문단 텍스트만 읽습니다. 표/이미지 해석은 별도 구현 전까지 제한
def _extract_docx(content: bytes) -> str:
    document = DocxDocument(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


# 이미지 자료 설명 텍스트
# 이미지는 도면 후보라는 사실만 corpus에 넣고, 구성요소/부호는 자동 확정하지 않음
def _image_material_text(filename: str, ocr_text: str) -> str:
    base = (
        f"이미지 자료 업로드: {filename}\n"
        "이 파일은 도면 또는 도안 후보로 업로드되었습니다. "
        "현재 시스템은 이미지의 구성요소와 부호를 자동 확정하지 않습니다. "
        "초안에는 이미지가 업로드되었다는 사실만 반영하고, 도면 명칭, 구성요소, 부호 설명은 사용자 확인이 필요합니다."
    )
    if ocr_text.strip():
        return f"{base}\n\n이미지 OCR 추출 텍스트:\n{ocr_text.strip()}"
    return base


# 이미지 OCR
def _extract_image(path: Path) -> str:
    try:
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=True):
            completed = subprocess.run(
                ["tesseract", str(path), "stdout", "-l", "kor+eng", "--psm", "6"],
                check=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        return completed.stdout.strip()
    except FileNotFoundError:
        return ""
